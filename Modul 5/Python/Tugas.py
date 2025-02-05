from docx import Document
from docx.shared import Inches
import os


def load_from_docx():
    pesanan_list = []
    if os.path.exists("pesanan_restoran.docx"):
        doc = Document("pesanan_restoran.docx")
        pesanan = {}
        for para in doc.paragraphs:
            text = para.text.strip()
            if text.startswith("ID Pesanan:"):
                if pesanan:
                    pesanan_list.append(pesanan)
                pesanan = {"id": int(text.split(":")[1].strip())}
            elif text.startswith("Nama Pelanggan:"):
                pesanan["nama"] = text.split(":")[1].strip()
            elif text.startswith("Menu Pesanan:"):
                pesanan["menu"] = text.split(":")[1].strip()
            elif text.startswith("Jumlah:"):
                pesanan["jumlah"] = int(text.split(":")[1].strip())
            elif text.startswith("Status:"):
                pesanan["status"] = text.split(":")[1].strip()
            elif text.startswith("Image Path:"):
                pesanan["image"] = text.split(":")[1].strip()
        if pesanan:
            pesanan_list.append(pesanan)
    return pesanan_list


def save_to_docx(pesanan_list):
    current_dir = os.path.dirname(os.path.abspath(__file__))
    file_path = os.path.join(current_dir, "pesanan_restoran.docx")

    doc = Document()
    doc.add_heading('Manajemen Pesanan Restoran', 0)

    for pesanan in pesanan_list:
        doc.add_paragraph(f"ID Pesanan: {pesanan['id']}")
        doc.add_paragraph(f"Nama Pelanggan: {pesanan['nama']}")
        doc.add_paragraph(f"Menu Pesanan: {pesanan['menu']}")
        doc.add_paragraph(f"Jumlah: {pesanan['jumlah']}")
        doc.add_paragraph(f"Status: {pesanan['status']}")
        
        
        if pesanan['image'] and os.path.exists(pesanan['image']):
            doc.add_paragraph(f"Image Path: {pesanan['image']}")
            try:
                doc.add_picture(pesanan['image'], width=Inches(3))  
            except Exception as e:
                doc.add_paragraph(f"Gagal menambahkan gambar: {e}")
        else:
            doc.add_paragraph("Tidak ada gambar")
        
        doc.add_paragraph("-" * 40)

    doc.save(file_path)


def create_pesanan(pesanan_list):
    id_pesanan = len(pesanan_list) + 1
    nama = input("Nama Pelanggan: ")
    menu = input("Menu Pesanan: ")
    jumlah = int(input("Jumlah: "))
    image_path = input("Path Image: ")
    
    print("\nPilih Status Pesanan:")
    print("1. Diproses")
    print("2. Selesai")
    status_pilihan = input("Masukkan pilihan (1/2): ")
    
    
    if status_pilihan == '2':
        status = "selesai"
    else:
        status = "diproses"

    pesanan = {
        "id": id_pesanan,
        "nama": nama,
        "menu": menu,
        "jumlah": jumlah,
        "status": status,  
        "image": image_path
    }

    pesanan_list.append(pesanan)
    save_to_docx(pesanan_list)
    print("Pesanan berhasil dibuat!")


def read_pesanan(pesanan_list):
    if len(pesanan_list) == 0:
        print("Tidak ada pesanan.")
    else:
        for pesanan in pesanan_list:
            print(f"ID Pesanan: {pesanan['id']}")
            print(f"Nama Pelanggan: {pesanan['nama']}")
            print(f"Menu Pesanan: {pesanan['menu']}")
            print(f"Jumlah: {pesanan['jumlah']}")
            print(f"Status: {pesanan['status']}")
            print(f"Image Path: {pesanan['image']}")
            print("-" * 40)


def update_pesanan(pesanan_list):
    id_pesanan = int(input("Masukkan ID Pesanan yang ingin diubah: "))
    for pesanan in pesanan_list:
        if pesanan['id'] == id_pesanan:
            print(f"Pesanan {pesanan['id']} ditemukan.")
            new_status = input(
                "Masukkan status baru (diproses/selesai/batal): ")
            pesanan['status'] = new_status
            save_to_docx(pesanan_list)
            print(f"Status pesanan {
                  id_pesanan} telah diubah menjadi {new_status}.")
            return
    print("Pesanan tidak ditemukan.")


def delete_pesanan(pesanan_list):
    id_pesanan = int(input("Masukkan ID Pesanan yang ingin dihapus: "))
    for pesanan in pesanan_list:
        if pesanan['id'] == id_pesanan:
            if pesanan['status'] == 'batal':
                pesanan_list.remove(pesanan)
                save_to_docx(pesanan_list)
                print(f"Pesanan {id_pesanan} telah dihapus.")
                return
            else:
                print("Pesanan hanya bisa dihapus jika statusnya 'batal'.")
                return
    print("Pesanan tidak ditemukan.")


def search_pesanan(pesanan_list):
    nama_cari = input("Masukkan nama pelanggan yang ingin dicari: ")
    ditemukan = False
    for pesanan in pesanan_list:
        if nama_cari.lower() in pesanan['nama'].lower():
            print(f"ID Pesanan: {pesanan['id']}")
            print(f"Nama Pelanggan: {pesanan['nama']}")
            print(f"Menu Pesanan: {pesanan['menu']}")
            print(f"Jumlah: {pesanan['jumlah']}")
            print(f"Status: {pesanan['status']}")
            print(f"Image Path: {pesanan['image']}")
            print("-" * 40)
            ditemukan = True
    if not ditemukan:
        print("Pesanan tidak ditemukan.")


def main():
    pesanan_list = load_from_docx()

    while True:
        print("\n1. Buat Pesanan")
        print("2. Baca Pesanan")
        print("3. Update Pesanan")
        print("4. Hapus Pesanan")
        print("5. Cari Pesanan")
        print("6. Keluar")

        pilihan = input("Pilih menu (1./2./3./4./5./6.): ")

        if pilihan == '1':
            create_pesanan(pesanan_list)
        elif pilihan == '2':
            read_pesanan(pesanan_list)
        elif pilihan == '3':
            update_pesanan(pesanan_list)
        elif pilihan == '4':
            delete_pesanan(pesanan_list)
        elif pilihan == '5':
            search_pesanan(pesanan_list)
        elif pilihan == '6':
            print("Terima kasih!")
            break
        else:
            print("Pilihan tidak valid, coba lagi.")


if __name__ == "__main__":
    main()